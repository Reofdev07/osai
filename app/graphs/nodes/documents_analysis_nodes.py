# import magic
import puremagic
import fitz
import os
import json
import re
import asyncio
import sqlite3
from datetime import datetime, timedelta

# --- Imports de librerías ---
from langchain_community.document_loaders import PyMuPDFLoader
from google.cloud import vision
from llama_parse import LlamaParse
from llama_index.core import SimpleDirectoryReader
from tenacity import retry, stop_after_attempt, wait_exponential

# --- Imports de tu propio proyecto ---
from app.schemas.graph_state import DocumentState
from app.utils.token_counter import count_tokens, update_usage_metadata
from app.utils.toon_helper import get_toon_context
from app.core.llm import create_llm
from app.core.config import settings
from app.core.database import DB_FILE, get_db_connection  # Importa la ruta y el getter
from app.schemas.agent_schemas import (
    ExtractionSummary, IntentAnalysis, SentimentUrgency, 
    ClassificationOutput, EntitiesOutput, PriorityOutput, ComplianceOutput, TagsOutput,
    MasterEnrichmentOutput, MegaEnrichmentOutput
)

# --- CONFIGURACIÓN GLOBAL PARA LOS NODOS ---
LLAMA_PARSE_FREE_LIMIT_WEEKLY = 7000
CONTEXT_WINDOW_LIMIT = 300000 # ~75k tokens, suficiente para documentos largos
llm = create_llm()
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = settings.GOOGLE_APPLICATION_CREDENTIALS

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=10),
    reraise=True
)
async def _invoke_llm_with_retry(runnable, prompt):
    """Wrapper con reintentos para llamadas al LLM."""
    return await runnable.ainvoke(prompt)

# === NODO 1: RUTA DE ENTRADA ===
async def analyze_and_route_node(state: DocumentState) -> DocumentState:
    """Analiza el tipo de archivo y decide la ruta inicial: texto, OCR o no soportado."""
    print("--- Nodo: Analizando tipo de archivo ---")
    file_path = state["file_path"]
    
    # Umbrales mejorados
    TEXT_RATIO_THRESHOLD = 100 # Promedio de caracteres por página para ser "pdf_text"
    MAX_PAGES_TO_SAMPLE = 3
    
    mime_type = puremagic.from_file(file_path, mime=True)
    
    if "pdf" in mime_type:
        try:
            with fitz.open(file_path) as doc:
                page_count = doc.page_count
                if page_count == 0: 
                    return {"file_type": "unsupported", "error": "PDF sin páginas"}
                
                # Analizamos una muestra (primeras 3 páginas) para decidir
                sample_text_length = 0
                pages_to_check = min(page_count, MAX_PAGES_TO_SAMPLE)
                
                for i in range(pages_to_check):
                    page = doc.load_page(i)
                    sample_text_length += len(page.get_text().strip())
                
                avg_text = sample_text_length / pages_to_check
                
                if avg_text > TEXT_RATIO_THRESHOLD:
                    print(f"--- Decisión: PDF Nativo (Promedio texto: {avg_text:.2f}) ---")
                    return {"file_type": "pdf_text", "page_count": page_count}
                else:
                    print(f"--- Decisión: PDF Escaneado (Promedio texto: {avg_text:.2f}) ---")
                    return {"file_type": "pdf_scanned", "page_count": page_count}
                    
        except Exception as e:
            print(f"Error analizando PDF: {e}")
            return {"file_type": "pdf_scanned"}
    elif "image" in mime_type:
        return {"file_type": "image", "page_count": 1}
    else:
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext in ['.docx', '.doc', '.xlsx', '.xls', '.csv', '.txt'] or any(t in mime_type for t in ["officedocument", "msword", "ms-excel", "csv", "text/plain"]):
            return {"file_type": "office_document", "page_count": 1}
        return {"file_type": "unsupported"}

# === NODOS DE EXTRACCIÓN DE TEXTO ===

# Ruta Barata: PDF con texto nativo
async def extract_from_text_pdf_node(state: DocumentState) -> DocumentState:
    """Extrae texto directamente de un PDF nativo usando PyMuPDF."""
    print("--- Worker: Extrayendo texto de PDF nativo (ruta barata) ---")
    try:
        loader = PyMuPDFLoader(state["file_path"])
        docs = loader.load()
        content = "".join([doc.page_content for doc in docs])
        page_count = len(docs)
        token_count = count_tokens(content)
        return {
            "raw_text": content, 
            "page_count": page_count, 
            "token_count": token_count,
            "extraction_method": "native_pdf",
            "extraction_pages": page_count
        }
    except Exception as e:
        return {"error": f"Error extrayendo texto de PDF: {e}"}

# Opción OCR 1: Google Vision
async def extract_with_google_vision_node(state: DocumentState) -> DocumentState:
    """
    Realiza OCR usando la API de Google Cloud Vision con optimización para documentos.
    Usa 'document_text_detection' que es superior para documentos con párrafos y tablas.
    """
    print("--- Worker (Expert): Realizando OCR con Google Vision (Document Analysis) ---")
    file_path = state["file_path"]
    job_id = state.get("job_id", "N/A")
    
    # Verificar credenciales antes de instanciar el cliente
    creds_path = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS")
    if not creds_path or not os.path.exists(creds_path):
        error_msg = f"Credenciales de GCP no encontradas en {creds_path}. Configure GOOGLE_APPLICATION_CREDENTIALS correctamente."
        print(f"Job [{job_id}]: {error_msg}")
        return {"error": error_msg, "extraction_pages": 0}

    client = vision.ImageAnnotatorClient()
    
    try:
        mime_type = puremagic.from_file(file_path, mime=True)
        extracted_content = ""
        page_count = 0
        
        if "pdf" in mime_type:
            all_text_pages = []
            with fitz.open(file_path) as doc:
                page_count = doc.page_count
                if page_count == 0: 
                    return {"error": "El PDF para OCR está vacío.", "extraction_pages": 0}
                
                for i, page in enumerate(doc):
                    print(f"Job [{job_id}]: Procesando página {i+1}/{page_count} con Google Vision...")
                    # Renderizado a alta calidad para mejor OCR
                    pix = page.get_pixmap(dpi=300)
                    image_bytes = pix.tobytes("png")
                    image = vision.Image(content=image_bytes)
                    
                    # 'document_text_detection' es mejor para documentos que 'text_detection'
                    response = client.document_text_detection(image=image)
                    
                    if response.error.message:
                        raise Exception(f"API Error pág {i+1}: {response.error.message}")
                    
                    if response.full_text_annotation:
                        all_text_pages.append(response.full_text_annotation.text)
            
            extracted_content = "\n\n--- Nueva Página ---\n\n".join(all_text_pages)
            
        elif "image" in mime_type:
            page_count = 1
            print(f"Job [{job_id}]: Procesando archivo de imagen con Google Vision...")
            with open(file_path, "rb") as image_file:
                content_bytes = image_file.read()
            
            image = vision.Image(content=content_bytes)
            response = client.document_text_detection(image=image)
            
            if response.error.message:
                raise Exception(f"API Error: {response.error.message}")
            
            if response.full_text_annotation:
                extracted_content = response.full_text_annotation.text
        else:
            return {"error": f"Tipo de archivo no soportado para Google Vision: {mime_type}"}
        
        token_count = count_tokens(extracted_content)
        print(f"Job [{job_id}]: Extracción con Google Vision finalizada. Páginas: {page_count}")
        
        return {
            "raw_text": extracted_content, 
            "page_count": page_count, 
            "token_count": token_count, 
            "error": None,
            "extraction_method": "google_vision",
            "extraction_pages": page_count
        }
    except Exception as e:
        error_message = f"Error crítico en Google Vision: {e}"
        print(f"Job [{job_id}]: {error_message}")
        return {"error": error_message, "extraction_pages": page_count or 0}

# === NUEVO: Extracción nativa de documentos Office (DOCX, XLSX, CSV, TXT) ===
async def extract_office_document_node(state: DocumentState) -> DocumentState:
    """
    Extrae texto plano de documentos de Office usando librerías nativas de Python.
    Soporta: .docx, .doc, .xlsx, .xls, .csv, .txt
    Similar a como ChatGPT lee estos archivos: directamente, sin OCR ni servicios externos.
    """
    print("--- Worker: Extrayendo texto de documento Office (nativo Python) ---")
    file_path = state["file_path"]
    job_id = state.get("job_id", "N/A")
    
    try:
        import pandas as pd
        from docx import Document as DocxDocument
        
        # Detectar el tipo real del archivo por extensión Y por mime
        file_ext = os.path.splitext(file_path)[1].lower()
        mime_type = puremagic.from_file(file_path, mime=True)
        
        extracted_content = ""
        page_count = 1
        
        # --- DOCX ---
        if file_ext in ['.docx'] or 'officedocument.wordprocessingml' in mime_type:
            print(f"Job [{job_id}]: Leyendo archivo DOCX con python-docx...")
            doc = DocxDocument(file_path)
            
            paragraphs_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs_text.append(text)
            
            # También extraer tablas del documento
            for i, table in enumerate(doc.tables):
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(row_data))
                if table_rows:
                    paragraphs_text.append(f"\n--- Tabla {i+1} ---")
                    paragraphs_text.extend(table_rows)
            
            extracted_content = "\n".join(paragraphs_text)
            # Estimar páginas (~3000 chars por página)
            page_count = max(1, len(extracted_content) // 3000)
        
        # --- XLSX / XLS ---
        elif file_ext in ['.xlsx', '.xls'] or any(t in mime_type for t in ['spreadsheetml', 'ms-excel']):
            print(f"Job [{job_id}]: Leyendo archivo Excel con pandas + openpyxl...")
            
            all_sheets_text = []
            xls = pd.ExcelFile(file_path)
            
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                df = _clean_dataframe(df)
                if not df.empty:
                    all_sheets_text.append(f"--- Hoja: {sheet_name} ({len(df)} filas, {len(df.columns)} columnas) ---")
                    all_sheets_text.append(df.to_markdown(index=False))
            
            extracted_content = "\n\n".join(all_sheets_text)
            page_count = len(xls.sheet_names)
        
        # --- CSV ---
        elif file_ext in ['.csv'] or 'csv' in mime_type:
            print(f"Job [{job_id}]: Leyendo archivo CSV con pandas...")
            df = pd.read_csv(file_path)
            df = _clean_dataframe(df)
            extracted_content = df.to_markdown(index=False)
        
        # --- TXT / Texto plano ---
        elif file_ext in ['.txt'] or 'plain' in mime_type:
            print(f"Job [{job_id}]: Leyendo archivo TXT como texto plano...")
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                extracted_content = f.read()
        
        # --- DOC (formato antiguo) ---
        elif file_ext in ['.doc'] or 'msword' in mime_type:
            print(f"Job [{job_id}]: Leyendo archivo DOC (formato legacy)...")
            # DOC antiguo es binario, intentamos extraer con antiword o lectura cruda
            try:
                import subprocess
                result = subprocess.run(['antiword', file_path], capture_output=True, text=True, timeout=30)
                if result.returncode == 0:
                    extracted_content = result.stdout
                else:
                    # Fallback: leer bytes y extraer lo que parezca texto
                    with open(file_path, "rb") as f:
                        raw = f.read()
                    extracted_content = raw.decode("utf-8", errors="ignore")
                    # Limpiar caracteres de control
                    extracted_content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', extracted_content)
            except Exception as doc_err:
                print(f"Job [{job_id}]: Error leyendo DOC legacy: {doc_err}")
                with open(file_path, "rb") as f:
                    raw = f.read()
                extracted_content = raw.decode("utf-8", errors="ignore")
                extracted_content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', extracted_content)
        
        else:
            return {"error": f"Formato de Office no reconocido: ext={file_ext}, mime={mime_type}"}
        
        # Validar que obtuvimos contenido
        if not extracted_content or not extracted_content.strip():
            return {"error": f"No se pudo extraer texto del archivo Office ({file_ext}). El archivo podría estar vacío o protegido."}
        
        token_count = count_tokens(extracted_content)
        print(f"Job [{job_id}]: Extracción Office finalizada. Caracteres: {len(extracted_content)}, Tokens: {token_count}")
        
        return {
            "raw_text": extracted_content,
            "page_count": page_count,
            "token_count": token_count,
            "error": None,
            "extraction_method": "native_office",
            "extraction_pages": page_count
        }
    
    except Exception as e:
        error_message = f"Error extrayendo documento Office: {e}"
        print(f"Job [{job_id}]: {error_message}")
        import traceback
        traceback.print_exc()
        return {"error": error_message, "extraction_pages": 0}

# Opción OCR 2: LlamaParse
async def extract_with_llama_parse_node(state: DocumentState) -> DocumentState:
    """
    Realiza OCR y parsing usando el método asíncrono nativo aload_data().
    """
    print("--- Worker: Extrayendo con LlamaParse (Modo Asíncrono Nativo) ---")
    file_path = state["file_path"]
    job_id = state.get("job_id", "N/A")

    try:
        # 1. Configurar el parser
        parser = LlamaParse(
            api_key=settings.LLAMA_CLOUD_API_KEY,
            result_type="markdown",
            verbose=True
        )

        # 2. Llamada asíncrona nativa (más eficiente que executors)
        print(f"Job [{job_id}]: Llamando a parser.aload_data() para: {file_path}")
        parsed_docs = await parser.aload_data([file_path])

        # 3. Procesar el resultado
        if not parsed_docs:
            return {"error": "LlamaParse no devolvió ningún documento."}

        # Unimos el contenido
        extracted_content = "\n\n".join([doc.get_content() for doc in parsed_docs])
        
        # Métricas
        page_count = parsed_docs[0].metadata.get('total_pages', 1) if parsed_docs[0].metadata else state.get("page_count_for_decision", 1)
        token_count = count_tokens(extracted_content)

        print(f"Job [{job_id}]: Extracción finalizada. Páginas: {page_count}, Tokens: {token_count}")

        return {
            "raw_text": extracted_content,
            "page_count": page_count,
            "token_count": token_count,
            "error": None,
            "extraction_method": "llama_parse",
            "extraction_pages": page_count
        }

    except Exception as e:
        error_message = f"Error en LlamaParse asíncrono: {e}"
        print(f"Job [{job_id}]: {error_message}")
        print(f"Job [{job_id}]: Iniciando Fallback a Google Vision...")
        # Fallback a Google Vision
        return await extract_with_google_vision_node(state)

# === NODOS DEL ORQUESTADOR ADAPTATIVO ===

async def count_pages_node(state: DocumentState) -> DocumentState:
    """Nodo rápido que cuenta las páginas de un documento para la decisión."""
    print("--- Orquestador: Contando páginas del documento ---")
    file_path = state["file_path"]
    job_id = state.get("job_id", "N/A")
    page_count = 1

    try:
        print(f"Job [{job_id}]: Contando páginas para el archivo: {file_path}")

        if not os.path.exists(file_path):
            error_msg = f"El archivo no existe en la ruta: {file_path}"
            print(f"Job [{job_id}]: ERROR - {error_msg}")
            return {"error": error_msg}

        mime_type = ""
        try:
            # La llamada de importación correcta es con guion bajo
            mime_type = puremagic.from_file(file_path, mime=True) 
        except Exception as e:
            error_msg = f"Error con la librería 'puremagic': {e}"
            print(f"Job [{job_id}]: ERROR - {error_msg}")
            return {"error": error_msg}

        if "pdf" in mime_type:
            try:
                with fitz.open(file_path) as doc:
                    page_count = doc.page_count
            except Exception as e:
                error_msg = f"Error con 'fitz' al abrir el PDF: {e}"
                print(f"Job [{job_id}]: ERROR - {error_msg}")
                return {"error": error_msg}
        
    except Exception as e:
        error_msg = f"Error inesperado en count_pages_node: {e}"
        print(f"Job [{job_id}]: ERROR - {error_msg}")
        return {"error": error_msg}

    print(f"--- Orquestador: Documento tiene {page_count} páginas ---")
    return {"page_count_for_decision": page_count}



async def adaptive_ocr_orchestrator_node(state: DocumentState) -> DocumentState:
    """Decide qué proveedor usar basándose en el contador de SQLite."""
    print("--- Orquestador Adaptativo: Tomando decisión de OCR ---")
    pages_to_process = state.get("page_count_for_decision", 1)
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = 'llama_parse_reset_timestamp'")
        last_reset_time = datetime.fromisoformat(cursor.fetchone()[0])
        if (datetime.now() - last_reset_time) >= timedelta(days=7):
            print("--- Orquestador (SQLite): Reseteando contador semanal ---")
            cursor.execute("UPDATE settings SET value = '0' WHERE key = 'llama_parse_usage'")
            cursor.execute("UPDATE settings SET value = ? WHERE key = 'llama_parse_reset_timestamp'", (datetime.now().isoformat(),))
            conn.commit()
        cursor.execute("SELECT value FROM settings WHERE key = 'llama_parse_usage'")
        current_usage = int(cursor.fetchone()[0])
    if (current_usage + pages_to_process) <= LLAMA_PARSE_FREE_LIMIT_WEEKLY:
        provider_choice = "llama_parse"
        print(f"Límite LlamaParse OK ({current_usage + pages_to_process}/{LLAMA_PARSE_FREE_LIMIT_WEEKLY}). Seleccionando LlamaParse.")
    else:
        provider_choice = "google_vision"
        print(f"Límite LlamaParse excedido ({current_usage}/{LLAMA_PARSE_FREE_LIMIT_WEEKLY}). Usando Google Vision.")
    return {"ocr_provider": provider_choice}

async def update_llama_parse_usage_node(state: DocumentState) -> DocumentState:
    """Si se usó un servicio de OCR/Parser, actualiza el contador en SQLite de forma atómica."""
    pages_processed = state.get("extraction_pages", 0)
    method = state.get("extraction_method")
    
    if pages_processed > 0:
        with get_db_connection() as conn:
            cursor = conn.cursor()
            if method == "llama_parse":
                cursor.execute("UPDATE settings SET value = CAST(value AS INTEGER) + ? WHERE key = 'llama_parse_usage'", (pages_processed,))
                key_name = "llama_parse_usage"
            elif method == "google_vision":
                cursor.execute("UPDATE settings SET value = CAST(value AS INTEGER) + ? WHERE key = 'google_vision_usage'", (pages_processed,))
                key_name = "google_vision_usage"
            else:
                return {} # No actualizamos para native_pdf en DB por ahora
                
            conn.commit()
            cursor.execute(f"SELECT value FROM settings WHERE key = ?", (key_name,))
            new_total = cursor.fetchone()[0]
            print(f"--- Orquestador (SQLite): Contador de {method} actualizado a {new_total} ---")
            
    return {}

# === Nodo Final para Archivos no Soportados ===
async def unsupported_file_node(state: DocumentState) -> DocumentState:
    print("--- Nodo: Archivo no soportado, finalizando flujo. ---")
    return {"error": "Tipo de archivo no soportado."}

# === NODOS DE ANÁLISIS DE CONTENIDO (TUS PROMPTS ORIGINALES) ===

async def summarize_and_get_subject_node(state: DocumentState) -> DocumentState:
    print("--- Worker (Expert): Resumo, Asunto y Fecha ---")
    if not state.get("raw_text", "").strip(): return {"error": "No hay texto para analizar."}
    
    prompt = f"""
    Act as a professional archivist expert in Colombian documents.
    1. Summarize the following document accurately in Spanish (3-4 sentences).
    2. Extract a clear, formal subject title.
    3. Identify the document date (the date of issuance as written in the text, e.g., 'Octubre 2025' or 'Octubre 1 de 2025'). 
    
    TEXT:
    {state['raw_text'][:CONTEXT_WINDOW_LIMIT]}
    """
    try:
        # Usamos include_raw=True para capturar la metadata de tokens
        print(f"📡 Enviando prompt al LLM...")
        runnable = llm.with_structured_output(ExtractionSummary, include_raw=True)
        result = await _invoke_llm_with_retry(runnable, prompt)
        print(f"✅ Respuesta recibida del LLM.")
        
        data = result['parsed']
        usage = result['raw'].usage_metadata
        
        return {
            "summary": data.resumen, 
            "subject": data.asunto,
            "document_date": data.fecha,
            "usage_metadata": usage # El reducer en el State se encarga de sumar
        }
    except Exception as e:
        print(f"!!! Error en summarize_and_get_subject_node: {e}")
        return {"errors": [f"Error en resumen: {e}"]}

async def mega_analysis_node(state: DocumentState) -> DocumentState:
    """
    Consolida Master Enrichment, Entidades, Prioridad y Compliance en UNA sola llamada.
    Reduce tokens de entrada en ~75% y evita errores de bloqueos 429 por Fan-Out concurrente.
    """
    print("--- Worker (Expert Mega): Análisis Global Unificado (Sustituyendo Fan-Out) ---")
    
    # Seguridad: Si no hay texto, no tiene sentido llamar al LLM
    raw_text = state.get("raw_text", "").strip()
    if not raw_text:
        return {"errors": ["No hay texto para realizar el mega-análisis."]}
        
    ctx = get_toon_context(state)
    
    prompt = f"""
    Act as an expert archival analyst and Legal Advisor specialized in Colombian documentation (Ley 594/2000, Ley 1755/2015, Acuerdo 060 AGN).
    Perform a multi-dimensional analysis of the document context provided.
    
    CONTEXT (TOON): {ctx}
    TEXT SEGMENT: {raw_text[:CONTEXT_WINDOW_LIMIT]}
    
    SCIENTIFIC/COGNITIVE TASKS:
    
    1. INTENT & SENTIMENT (Intencion y Sentimiento):
       - Intent Options: Solicitar Información, Presentar Queja/Reclamo, Radicar para Pago, 
         Entregar Documentación, Iniciar Trámite Nuevo, Notificar Decisión, Consulta General, Informativo/Cortesía.
       - Tone (-1 to 1) and urgency (Baja, Media, Alta, Crítica).
    
    2. CLASSIFICATION & TAGS (Clasificacion y Etiquetas):
       - Typology Options: Acto Administrativo, Contrato, Informe, Factura, Historia Laboral, Hoja de Vida, 
         Solicitud (PQRS), Tutela, Comunicación Oficial, Certificado, Otro.
       - Create 5-7 relevant Spanish tags.
       
    3. ENTITIES (Entidades):
       - Identify People, Organizations, Dates, Amounts, and specific Codes (Radicados).
       - Map relevant Facts and build a Timeline.
       - Empty list [] if not found. Do NOT invent data.
       
    4. LEGAL PRIORITY (Prioridad Legal Ley 1755):
       - Docs/Info: 10 days. Queries: 30 days. General/PQRS: 15 days.
       - MUST cite the specific article/inciso.
       
    5. COMPLIANCE (Conformidad Archivística):
       - Check Sender Identification, Recipient correctness, Purpose clarity, Signature presence.
       - Summarize compliance and provide detailed recommendations.
       
    Your output MUST be highly professional and perfectly structured. Do not invent data. Return empty structures if necessary.
    """
    
    try:
        runnable = llm.with_structured_output(MegaEnrichmentOutput, include_raw=True)
        result = await _invoke_llm_with_retry(runnable, prompt)
        
        data = result['parsed']
        usage = result['raw'].usage_metadata
        
        return {
            "intent_analysis": data.intencion.dict(),
            "sentiment_analysis": {
                "sentimiento": {
                    "etiqueta": data.sentimiento_urgencia.etiqueta,
                    "puntuacion": data.sentimiento_urgencia.puntuacion,
                    "justificacion": data.sentimiento_urgencia.justificacion
                },
                "urgencia": {
                    "nivel": data.sentimiento_urgencia.urgencia_nivel,
                    "justificacion": data.sentimiento_urgencia.urgencia_justificacion
                }
            },
            "classification": {
                "tipologia_documental": data.clasificacion.tipologia_documental, 
                "confianza": data.clasificacion.confianza
            },
            "tags": data.etiquetas,
            "entities": data.entidades.dict(),
            "priority_analysis": data.prioridad.dict(),
            "compliance_analysis": {
                "cumple_normativa": data.conformidad.cumple_normativa,
                "resumen": data.conformidad.resumen_ejecutivo,
                "detalles": data.conformidad.analisis_detallado
            },
            "usage_metadata": usage
        }
    except Exception as e:
        print(f"!!! Error en mega_analysis_node: {e}")
        return {"errors": [f"Error en MEGA análisis estructurado: {e}"]}
